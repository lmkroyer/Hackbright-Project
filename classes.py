"""FACTORY OF FORM OBJECTS."""

from textblob import TextBlob
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from defenses import all_defenses
import inflect


# FIXME: remove all punctuation!!!! so don't have to do weird searches

class Complaint(TextBlob):
    def __init__(self, decoded_text):
        super(Complaint, self).__init__(decoded_text)

        self.word_list = self.words
        self.sentence_list = self.sentences
        self.nouns = self.noun_phrases
        self.plaintiff_fname = self.get_plaintiff_fname()
        self.plaintiff_lname = self.get_plaintiff_lname()
        self.case_no = self.get_case_no()
        self.county = self.get_county()
        self.state = self.get_state()
        self.defendant_fname = self.get_defendant_fname()
        self.defendant_lname = self.get_defendant_lname()
        self.defendant_residence = self.get_defendant_residence()
        self.amount_claimed = self.get_amount_claimed()
        self.claim = self.get_claim()
        self.counsel_fname = self.get_counsel_fname()
        self.counsel_lname = self.get_counsel_lname()
        self.counsel_firm = self. get_counsel_firm()
        self.complaint_date = self.get_complaint_date()
        self.legal_basis = self.get_legal_basis()
        # self.injury_date = self.get_injury_date()
        # self.injury_location = self.get_injury_location()
        # self.injury_description = self.get_injury_description()

        # etc...

# FIXME: refactor option to consolidate double ifs in one line, and return the self. (no name)
# FIXME: remove punctuation from word list so can not hardcode that into checks
# FIXME: .correct() everything that comes out
# FIXME: a way to grab other defendants and/or check if company or person

# FIXME: get all names by making a noun phrase list!!!

    # Could also check if what comes after starts with upper case, and if so pull it
    def get_plaintiff_fname(self):
        """Return the plaintiff's first name."""

        # for i in range(len(self.word_list)):
        #     if self.word_list[i] == 'Plaintiff':
        #         plaintiff_fname = self.word_list[i - 4]
        #         return plaintiff_fname.capitalize()

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'plaintiff':
                full_name = self.nouns[i - 1]
                plaintiff_fname = full_name.split(' ')[0]
                return plaintiff_fname.capitalize()

    def get_plaintiff_lname(self):
        """Return the plaintiff's last name."""

        # for i in range(len(self.word_list)):
        #     if self.word_list[i] == 'Plaintiff':
        #         plaintiff_lname = self.word_list[i - 2]
        #         return plaintiff_lname.capitalize()
        for i in range(len(self.nouns)):
            if self.nouns[i] == 'plaintiff':
                full_name = self.nouns[i - 1]
                plaintiff_lname = full_name.split(' ')[1]
                return plaintiff_lname.capitalize()

    def get_case_no(self):
        """Return the plaintiff's last name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'No':
                case_no = self.word_list[i + 1]
                return case_no


    def get_county(self):
        """Return the county."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'COUNTY':
                if self.word_list[i + 1] == 'OF':
                    county = self.word_list[i + 2]
                    return county.capitalize()


    def get_state(self):
        """Return the state."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'STATE':
                if self.word_list[i + 1] == 'OF':
                    state = self.word_list[i + 2]
                    return state.capitalize()


    def get_defendant_fname(self):
        """Return the defendant's first name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'defendant':
                return self.word_list[i + 1]
        # for i in range(len(self.nouns)):
        #     if self.nouns[i] == 'defendant':
        #         defendant_fname = self.nouns[i + 1]
        #         return defendant_fname.capitalize()

    def get_defendant_lname(self):
        """Return the defendant's last name."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'defendant':
                return self.word_list[i + 2]
        # for i in range(len(self.nouns)):
        #     if self.nouns[i] == 'defendant':
        #         defendant_lname = self.nouns[i + 2]
        #         return defendant_lname.capitalize()

    def get_amount_claimed(self):
        """Return the dollar of damages requested."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'Amount':
                if self.word_list[i + 1] == 'claimed':
                    return self.word_list[i + 2]


    def get_claim(self):
        """Return the type of claim as either the default PI or an error message."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('Personal Injury'):
                return 'Personal Injury'
            else:
                return 'UNKNOWN'


    def get_legal_basis(self):
        """Return the types of legal claims made in the complaint.

        TODO: (2.0) allow multiple legal bases."""

        # result = []

        # for i in range(len(self.word_list)):

        # #     if self.sentence_list[i].find('Negligence') and 'Negligence' not in result:
        # #         result.append('Negligence')
        # #     # etc.
        # # result = ' '.join(result)
        #     if self.word_list[i] == 'Claim' and self.word_list[i + 1] == 'for' and self.word_list[i + 2] == 'Relief':
        #     #     if self.word_list[i + 3] not in result:
        #     #         result.append(self.word_list[i + 3])
        #     # result = ' '.join(result)
        #         return self.word_list[i + 3]

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'claim' and self.nouns[i + 1] == 'relief':
                result = self.nouns[i + 2]
                return result.capitalize()


    def get_defendant_residence(self):
        """Return the defendant's city, county, state of residence."""

        for i in range(len(self.word_list)):
            if self.word_list[i] == 'resides':
                defendant_residence = self.word_list[i + 2:i + 6]
                return ' '.join(defendant_residence)


    def get_counsel_fname(self):
        """Return the opposing counsel's first name."""

        # for i in range(len(self.word_list)):
        #     if self.word_list[i] == 'P.C':
        #         return self.word_list[i + 2]

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'llc':
                full_name = self.nouns[i + 1]
                counsel_fname = full_name.split(' ')[0]
                return counsel_fname.capitalize()


    def get_counsel_lname(self):
        """Return the opposing counsel's last name."""

        # for i in range(len(self.word_list)):
        #     if self.word_list[i] == 'P.C' or self.word_list[i] == 'LLC':
        #         return self.word_list[i + 4]
        for i in range(len(self.nouns)):
            if self.nouns[i] == 'llc':
                full_name = self.nouns[i + 1]
                counsel_lname = full_name.split(' ')[2]
                return counsel_lname.capitalize()


    def get_counsel_firm(self):
        """Return the opposing counsel's firm name."""

        # for i in range(len(self.sentence_list)):

        #     if self.sentence_list[i].find('LLC') or self.sentence_list[i].find('P.C'):
        #         return self.sentence_list[i]

        for i in range(len(self.nouns)):
            if self.nouns[i] == 'llc':
                firm = self.nouns[i -1]
                firm = firm.title()
                #FIXME: titlecase firm (can from titlecase import titlecase -- look in textblob)
                org = self.nouns[i]
                org = org.upper()
                # preposition.upper()
                return firm + ' ' + org
                # return full_firm


    def get_complaint_date(self):
        """Stampe the complaint with the date processed."""

        return datetime.utcnow()


    def get_injury_date(self):

        return 'N/A'

        # regex: 'On *** plaintiff was'


    def get_injury_location(self):

        return 'N/A'


    def get_injury_description(self):

        return 'N/A'

class Answer(object):

    plaintiff_fname = plaintiff_lname = defendant_fname = defendant_lname = None
    case_state = case_county = user_fname = user_lname = user_mailing_address = None
    user_email = user_firm_name = case_no = None

    def __init__(self, complaint, user, defenses):

# pass a complaint into the answer, and loop through complaint attributes and set to answer
# method that takes in self and generates a document
# make info inside brackets match so can set as variable and check all brackets
# look up: get attr and set attr (takes in a variable, find attribute that matches that and set)
# make a dictionary as a class attribute with keys as attribute name, value as right side
#   use setattr(self, "left side", )

        self.complaint = complaint
        self.user = user
        self.defenses = defenses
        self.plaintiff_fname = complaint.case.plaintiffs[0].fname
        self.plaintiff_lname = complaint.case.plaintiffs[0].lname
        self.defendant_fname = complaint.case.defendants[0].fname
        self.defendant_lname = complaint.case.defendants[0].lname
        self.case_county = complaint.case.county
        self.case_state = complaint.case.state
        self.user_fname = user.fname
        self.user_lname = user.lname
        self.user_email = user.email
        self.user_mailing_address = user.mailing_address
        self.user_firm_name = user.firm_name
        self.case_no = str(complaint.case.case_no)


    def insert_information(self):
        """Adds custom information into the answer template.

        Returns a modified docx file."""
        # FIXME: add pronoun awareness (textblob)!!! on the extract from dictionary point.
        # give user a change to choose female, male or plural
        # user textblob on paragraphs to insert to convert to plural or pronoun

        # Grab a value, turn to textblob object, change gender pronount, and then insert
        # FIXME: change civil code #s so that state gets input as variable

        # Make a list of all attributes on an Answer class
        attrs = [attr for attr in Answer.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(Answer.__dict__[attr]))]

        # Make a list of all upper case attributes
        cap_attrs = [attr.upper() for attr in Answer.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(Answer.__dict__[attr]))]

        # Make a Document object for Python-docx
        answer = Document('forms/answer_template.docx')
        style = answer.styles['Normal']
        font = style.font
        font.size = Pt(12)

        # font.highlight_color = WD_COLOR.YELLOW

        # Preserve format: replace lowercase tags with values
        for attr_name in attrs:

            for paragraph in answer.paragraphs:

                if attr_name in paragraph.text:
                    paragraph.text = (
                        paragraph.text.replace(attr_name,
                                               getattr(self, attr_name)))

        # Preserve format: replace uppercase tags with values
        for attr_name in cap_attrs:

            for paragraph in answer.paragraphs:

                if attr_name in paragraph.text:
                    attr_data = getattr(self, attr_name.lower())
                    attr_data = attr_data.upper()
                    paragraph.text = (
                        paragraph.text.replace(attr_name, attr_data))

        counter1 = 2
        counter2 = 3
        # FIXME: make this its own method on the class
        for defense in self.defenses:

            for p in answer.paragraphs:
                # Begin making paragraphs
                if '***' in p.text:
                    # Grab the legalese from the dictionary - a string of text
                    legalese = all_defenses[defense]

                    # Use inflect library to convert int to spelled digit
                    convert = inflect.engine()
                    spell_digit = convert.number_to_words(counter1)
                    spell_ordinal = convert.ordinal(spell_digit)
                    spell_ordinal = spell_ordinal.upper()

                    # Add the header for each paragraph
                    prior_paragraph = p.insert_paragraph_before()
                    prior_paragraph.add_run("{spell_ordinal} AFFIRMATIVE DEFENSE".format(spell_ordinal=spell_ordinal)).bold = True
                    prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    prior_paragraph = p.insert_paragraph_before()
                    counter1 += 1

                    # Grammar check
                    if spell_ordinal != 'eighth' or 'eleventh':
                        custom_intro = "As a {ordinal}, separate, and affirmative defense".format(spell_ordinal=spell_ordinal)
                    else:
                        custom_intro = "As an {ordinal}, separate, and affirmative defense".format(spell_ordinal=spell_ordinal)

                    # Add the meat of the legal text, with formatting
                    legalese = str(counter2) +'.' + '\t' + '\t' + custom_intro + legalese
                    prior_paragraph = p.insert_paragraph_before(legalese)
                    paragraph_format = prior_paragraph.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.25)
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    counter2 += 1


        # Check for paragraph titles to bold

        filename = 'answer_{case_no}.docx'.format(case_no=self.case_no)
        answer.save('filestorage/{filename}'.format(filename=filename))
        return filename

        # answer.save('filestorage/{case_no}.docx'.format(case_no=self.case_no))
        # return something to pass to display




# don't assume end coordinate
# give user option to draw
# function returns list
# iterate over and crop box/image for each item
# draw image with text returned (opencv)
# algorithm: given this location and image, where can i place it?
# tiling project







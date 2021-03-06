# coding: utf8

"""FACTORY OF CORPORATE FORM OBJECTS."""

from textblob import TextBlob
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
import inflect


class LPAForm(object):

    _fund_ = _fundState_ = _fundPpp_ = None
    _gp_ = _gpState_ = _gpAddress_ = _gpEmail_ = _gpOrgType_ = _gpSigParty_ = None
    _im_ = _imState_ = _imAddress_ = _imEmail_ = _imOrgType_ = None
    _sigDate_ = None

    def __init__(self, fund, fund_state, fund_ppp,
                 gp, gp_state, gp_address, gp_email, gp_sig_party,
                 im, im_state, im_address, im_email,
                 sig_date):

        self._fund_ = fund
        self._fundState_ = fund_state
        self._fundPpp_ = fund_ppp
        self._gp_ = gp
        self._gpOrgType_ = self.spell_gp_org()
        self._gpState_ = gp_state
        self._gpAddress_ = gp_address
        self._gpEmail_ = gp_email
        self._gpSigParty_ = gp_sig_party
        self._im_ = im
        self._imOrgType_ = self.spell_im_org()
        self._imState_ = im_state
        self._imAddress_ = im_address
        self._imEmail_ = im_email
        self._sigDate_ = sig_date


    def spell_gp_org(self):
        """Return the full spelling of the GP's organizational type."""

        gp = self._gp_
        gp_org = gp.split(' ')[-1]

        if gp_org == 'LP':
            return 'Limited Partnership'
        elif gp_org == 'LLC':
            return 'Limited Liability Company'


    def spell_im_org(self):
        """Return the full spelling of the IM's organizational type."""

        im = self._im_
        im_org = im.split(' ')[-1]

        if im_org == 'LP':
            return 'Limited Partnership'
        elif im_org == 'LLC':
            return 'Limited Liability Company'


    def insert_information(self):
        """Adds custom information into the request for production of documents template.

        Returns a modified docx file."""

        # Make a list of all attributes on a RequestProDocs class
        attrs = [attr for attr in LPAForm.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(LPAForm.__dict__[attr]))]

        # Make a list of all upper case attributes
        cap_attrs = [attr.upper() for attr in LPAForm.__dict__.keys()
                 if (not attr.startswith("__") and
                     not callable(LPAForm.__dict__[attr]))]

        # Make a Document object for Python-docx
        lpa = Document('forms/lpa_template.docx')
        style = lpa.styles['Normal']
        font = style.font
        font.size = Pt(10)

        # Preserve format: replace lowercase tags with values
        for attr_name in attrs:

            for paragraph in lpa.paragraphs:

                if attr_name in paragraph.text:
                    paragraph.text = (
                        paragraph.text.replace(attr_name,
                                               getattr(self, attr_name)))

        # Preserve format: replace uppercase tags with values
        for attr_name in cap_attrs:

            for paragraph in lpa.paragraphs:

                if attr_name in paragraph.text:
                    attr_data = getattr(self, attr_name.lower())
                    attr_data = attr_data.upper()
                    paragraph.text = (
                        paragraph.text.replace(attr_name, attr_data))

        # Make a new filename from fund name
        filename = 'LPA_{fund}.docx'.format(fund=self._fund_)
        # Save the modified document with the new filename
        lpa.save('filestorage/{filename}'.format(filename=filename))
        # Return filename to pass to display
        return filename


class PPMForm(TextBlob):

    def __init__(self, decoded_text):
        super(PPMForm, self).__init__(decoded_text)

        self.sentence_list = self.sentences
        self.fund = self.get_fund_name()
        self.mgmt_fee = self.get_mgmt_fee()
        self.jurisdiction = self.get_jurisdiction()
        self.manager = self.get_manager()
        self.principals = self.get_principals()
        self.removal = self.get_removal()
        self.leverage = self.get_indebtedness()
        self.min_commitment = self.get_min_commitment()
        self.transfers = self.get_transfers()
        self.reinvestment = self.get_reinvestment()


    def get_fund_name(self):
        """Return the name of the reviewing fund."""

        result = self.sentence_list[0]
        string = str(result)

        name = re.search(r'(?:(?!CONFIDENTIAL).)*', string).group()

        return name.upper()


    def get_mgmt_fee(self):
        """Return the management fee of the fund."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('“Management Fee”'.decode('utf-8')) != -1 and self.sentence_list[i].find('%') != -1:
                # Returns a sentence object
                sentence = self.sentence_list[i]
                # Get string from sentence object
                string = str(sentence)

                percent = re.search(r"\d+.\d+\%", string).group()
                payable_when = re.search(r"(?<=\bwhich will be payable\s)(?:[A-Za-z]+){2}[A-Za-z]+", string).group()
                fee = percent + ', payable ' + payable_when

                return fee


    def get_jurisdiction(self):
        """Return the jurisdiction of the fund entity."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('“Fund”'.decode('utf-8')) != -1 and self.sentence_list[i].find('under the laws of') != -1:

                sentence = self.sentence_list[i]
                string = str(sentence)
                return re.search(r"(?<=\bunder the laws of\s)(\w+)", string).group()


    def get_manager(self):
        """Return the manager of the fund ."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('will be the manager') != -1:

                sentence = self.sentence_list[i]
                string = str(sentence)

                return re.search(r"^.*(?=(\,))", string).group()


    def get_principals(self):
        """Return the principal(s) of the fund."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('principals of the Manager') != -1:
                sentence = self.sentence_list[i]
                string = str(sentence)
                list_words = string.split(' ')

                return ' '.join(list_words)


    def get_removal(self):
        """Return the removal provision."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('Manager may be removed') != -1:

                sentence = self.sentence_list[i]
                return str(sentence)


    def get_indebtedness(self):
        """Return the fund's allowed use of leverage."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('indebtedness') != -1 and self.sentence_list[i].find('may') != -1:
                result = self.sentence_list[i]

                return str(result)


    def get_min_commitment(self):
        """Return the minimum capital commitment."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('minimum capital commitment') != -1:
                sentence = self.sentence_list[i]
                string = str(sentence)
                num = re.search(r"\d+.\d+", string).group()

                return '$' + num


    def get_reinvestment(self):
        """Return the reinvestment provision."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('reinvestment') != -1 and self.sentence_list[i].find('subject') != -1:
                sentence = self.sentence_list[i]

                return str(sentence)


    def get_transfers(self):
        """Return the transfer permissions provision."""

        for i in range(len(self.sentence_list)):

            if self.sentence_list[i].find('proposed transfers') != -1:
                sentence = self.sentence_list[i]

                return str(sentence)